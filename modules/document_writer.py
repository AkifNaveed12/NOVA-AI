"""
MODULE — Document Writer
Takes generated assignment content and produces:
- PDF output via fpdf2
- DOCX output via python-docx
Saves to nova_outbox/ with timestamped filename.
"""

import os
from datetime import datetime
from pathlib import Path
from fpdf import FPDF
import docx as _docx


class DocumentWriter:

    def __init__(self, outbox_path: str = "nova_outbox"):
        self.outbox = Path(outbox_path)
        self.outbox.mkdir(exist_ok=True)

    def write(self, title: str, content: str, subject: str,
              output_format: str = "pdf") -> dict:
        """
        Write the assignment to disk.
        Returns: {"success": bool, "path": str, "filename": str, "format": str}
        """
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        safe_subject = "".join(c for c in subject if c.isalnum() or c in " _").strip().replace(" ", "_")
        filename_base = f"{timestamp}_{safe_subject}_Assignment"

        if output_format.lower() == "docx":
            return self._write_docx(title, content, filename_base)
        else:
            return self._write_pdf(title, content, filename_base)

    def _clean_latin1(self, text: str) -> str:
        replacements = {
            '\u2018': "'",
            '\u2019': "'",
            '\u201c': '"',
            '\u201d': '"',
            '\u2013': '-',
            '\u2014': '-',
            '\u2022': '*',
            '\u2026': '...',
        }
        for orig, repl in replacements.items():
            text = text.replace(orig, repl)
        return text.encode('latin-1', 'replace').decode('latin-1')

    def _write_pdf(self, title: str, content: str, filename_base: str) -> dict:
        try:
            title = self._clean_latin1(title)
            content = self._clean_latin1(content)
            pdf = FPDF()
            pdf.set_margins(20, 20, 20)
            pdf.add_page()

            # Title
            pdf.set_font("Helvetica", "B", 16)
            pdf.multi_cell(0, 10, title, align="C")
            pdf.ln(8)

            # Horizontal line
            pdf.set_draw_color(100, 100, 100)
            pdf.line(20, pdf.get_y(), 190, pdf.get_y())
            pdf.ln(6)

            # Body — parse headings (lines starting with # or ALL CAPS short lines)
            pdf.set_font("Helvetica", "", 11)
            for line in content.split("\n"):
                line = line.strip()
                if not line:
                    pdf.ln(4)
                    continue
                if line.startswith("## ") or line.startswith("# "):
                    pdf.set_font("Helvetica", "B", 13)
                    pdf.multi_cell(0, 8, line.lstrip("# ").strip())
                    pdf.set_font("Helvetica", "", 11)
                    pdf.ln(2)
                elif line.startswith("**") and line.endswith("**"):
                    pdf.set_font("Helvetica", "B", 11)
                    pdf.multi_cell(0, 7, line.strip("*"))
                    pdf.set_font("Helvetica", "", 11)
                else:
                    pdf.multi_cell(0, 7, line)

            filepath = str(self.outbox / f"{filename_base}.pdf")
            pdf.output(filepath)
            return {"success": True, "path": filepath,
                    "filename": f"{filename_base}.pdf", "format": "pdf"}
        except Exception as e:
            return {"success": False, "error": str(e), "path": "", "filename": ""}

    def _write_docx(self, title: str, content: str, filename_base: str) -> dict:
        try:
            doc = _docx.Document()
            doc.add_heading(title, 0)

            for line in content.split("\n"):
                line = line.strip()
                if not line:
                    doc.add_paragraph("")
                    continue
                if line.startswith("## "):
                    doc.add_heading(line[3:], level=2)
                elif line.startswith("# "):
                    doc.add_heading(line[2:], level=1)
                else:
                    doc.add_paragraph(line)

            filepath = str(self.outbox / f"{filename_base}.docx")
            doc.save(filepath)
            return {"success": True, "path": filepath,
                    "filename": f"{filename_base}.docx", "format": "docx"}
        except Exception as e:
            return {"success": False, "error": str(e), "path": "", "filename": ""}
